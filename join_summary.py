import docx
import pandas as pd


if __name__ == '__main__':
    # tmp/index.csvを読み込む
    df = pd.read_csv('tmp/index.csv')

    # docxを新規作成
    doc = docx.Document()

    for index, row in df.iterrows():
        # summaryから呼び込むためのファイルパスに変更する
        summary_path = row['file_name'].replace('prompt', 'summary')

        # summaryを読み込む
        with open(summary_path) as f:
            summary = f.read()
        
        # もし、titleが0文字だったら、見出しを追加しない
        if len(row['title']) > 0:
            # docxにtitleを追加する
            doc.add_heading(row['title'], level=1)

        # docxにsummaryを追加する
        doc.add_paragraph(summary)
    
    # docxを保存する
    doc.save('tmp/summary.docx')